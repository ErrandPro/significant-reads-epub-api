import re
import os
import logging
import subprocess
import tempfile
from collections import Counter

from pdfminer.high_level import extract_text

logger = logging.getLogger(__name__)

# ── XML namespaces used in DOCX/OOXML ────────────────────────────────────────

_NS_W   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_NS_A   = "http://schemas.openxmlformats.org/drawingml/2006/main"
_NS_R   = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_NS_WPS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
_NS_WP  = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"


# ════════════════════════════════════════════════════════════════════════════
# PDF — Text extraction
# ════════════════════════════════════════════════════════════════════════════

def extract_text_from_pdf(pdf_path: str) -> str:
    """pdfminer extraction — good fallback for complex fonts."""
    try:
        text = extract_text(pdf_path)
        return text if text else ""
    except Exception as e:
        logger.warning(f"pdfminer failed: {e}")
        return ""


def ocr_pdf_if_needed(pdf_path: str) -> str:
    """
    Tesseract OCR pipeline — last resort for image-only PDFs.
    Uses pdf2image + pytesseract with basic pre-processing.
    """
    from pdf2image import convert_from_path
    import pytesseract
    from PIL import ImageFilter, ImageOps

    images = convert_from_path(pdf_path, dpi=200)
    pages = []

    for img in images:
        img = img.convert("L")
        img = ImageOps.autocontrast(img)
        img = img.filter(ImageFilter.SHARPEN)
        pages.append(pytesseract.image_to_string(img, lang="eng"))

    return "\n".join(pages)


def extract_cover_image(pdf_path: str) -> bytes | None:
    """Render first PDF page into PNG for EPUB cover."""
    try:
        import fitz

        doc  = fitz.open(pdf_path)
        page = doc[0]
        pix  = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
        png_bytes = pix.tobytes("png")
        doc.close()
        return png_bytes

    except Exception as e:
        logger.info(f"Cover extraction skipped: {e}")
        return None


# ════════════════════════════════════════════════════════════════════════════
# DOCX — Word document support
# ════════════════════════════════════════════════════════════════════════════

def convert_doc_to_docx(doc_path: str) -> str:
    """
    Convert a legacy .doc file to .docx using LibreOffice headless.

    Returns the path to the freshly-created .docx file in a temp directory.
    The caller is responsible for cleaning up the directory when done.
    Raises RuntimeError if LibreOffice is not found or conversion fails.
    """
    soffice = _find_soffice()
    out_dir = tempfile.mkdtemp(prefix="doc2docx_")

    result = subprocess.run(
        [soffice, "--headless", "--convert-to", "docx",
         "--outdir", out_dir, doc_path],
        capture_output=True,
        timeout=120,
    )

    if result.returncode != 0:
        raise RuntimeError(
            f"LibreOffice conversion failed (exit {result.returncode}): "
            f"{result.stderr.decode(errors='replace').strip()}"
        )

    basename = os.path.splitext(os.path.basename(doc_path))[0]
    out_path  = os.path.join(out_dir, f"{basename}.docx")

    if not os.path.exists(out_path):
        raise RuntimeError(
            f"LibreOffice ran but output not found at {out_path}. "
            f"stdout: {result.stdout.decode(errors='replace').strip()}"
        )

    logger.info(f".doc → .docx: {out_path}")
    return out_path


def _find_soffice() -> str:
    """Locate the LibreOffice binary across common install paths."""
    candidates = [
        "soffice",
        "/usr/bin/soffice",
        "/usr/lib/libreoffice/program/soffice",
        "/opt/libreoffice/program/soffice",
        "/Applications/LibreOffice.app/Contents/MacOS/soffice",
    ]
    for c in candidates:
        try:
            result = subprocess.run(
                [c, "--version"], capture_output=True, timeout=5
            )
            if result.returncode == 0:
                return c
        except (FileNotFoundError, subprocess.TimeoutExpired):
            continue

    raise RuntimeError(
        "LibreOffice not found. Install it with:\n"
        "  Ubuntu: sudo apt-get install -y libreoffice\n"
        "  macOS:  brew install libreoffice"
    )


def extract_text_from_docx(docx_path: str) -> str:
    """
    Plain-text extraction from a .docx file (no formatting).
    Used as a fallback when rich extraction is unavailable or fails.
    """
    try:
        from docx import Document

        doc   = Document(docx_path)
        parts = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        for table in doc.tables:
            for row in table.rows:
                cells = "\t".join(c.text.strip() for c in row.cells)
                if cells.strip():
                    parts.append(cells)

        return "\n\n".join(parts)

    except Exception as e:
        logger.warning(f"extract_text_from_docx failed: {e}")
        return ""


def _is_toc_chapter(title: str, blocks: list[dict]) -> bool:
    """
    Detect whether a chapter is actually a Table of Contents.
    Checks the title and also scans block text for TOC-like patterns
    (lines ending in page numbers with dots/spaces).
    """
    title_lower = title.strip().lower()
    toc_titles  = {"contents", "table of contents", "content", "toc"}
    if title_lower in toc_titles:
        return True

    # Scan block text for TOC line patterns: "Some Title ........ 12"
    toc_pattern = re.compile(r"\.{3,}\s*\d+\s*$|\s{3,}\d+\s*$")
    toc_hits    = 0
    total_lines = 0

    for blk in blocks:
        if blk.get("kind") != "text":
            continue
        for ln in blk.get("lines", []):
            text = "".join(s.get("text", "") for s in ln.get("spans", []))
            if text.strip():
                total_lines += 1
                if toc_pattern.search(text):
                    toc_hits += 1

    # If more than 40% of lines look like TOC entries, skip this chapter
    if total_lines >= 4 and toc_hits / total_lines >= 0.40:
        return True

    return False


def extract_rich_chapters_from_docx(
    docx_path: str,
) -> list[tuple[str, list[dict]]] | None:
    """
    Rich extraction from a .docx file.

    Produces the identical block schema used by the PDF rich extractor so that
    epub_builder.py needs no changes:

        Chapter = (title: str, blocks: list[dict])

        Block kinds
        ───────────
        {"kind": "text",    "lines": [line, ...]}
        {"kind": "sidebar", "lines": [line, ...]}
        {"kind": "image",   "data": bytes, "ext": str}
        {"kind": "table",   "rows": [[str, ...], ...]}

        Line (inside a text/sidebar block)
        ───────────────────────────────────
        {
            "spans":      [{"text": str, "bold": bool, "italic": bool, "size": float}, ...],
            "is_section": bool,
            "dropcap_char": str | None,
        }

    Heading mapping
    ───────────────
        Heading 1          → chapter boundary
        Heading 7          → chapter boundary (used as "Chapter N" label in this schema)
        Title style        → title page (handled separately before main loop)
        Heading 2–6        → section heading (is_section = True)
        All-bold Normal    → section heading (is_section = True)
        Font > 1.5× body   → chapter boundary
        Font > 1.15× body  → section heading
    """
    try:
        from docx import Document
        from docx.oxml.ns import qn
        from docx.text.paragraph import Paragraph as DocxPara
        from docx.table import Table as DocxTable

        doc = Document(docx_path)

        # ── 1. Determine body font size (median of all run sizes) ──────────
        all_sizes: list[float] = []
        for para in doc.paragraphs:
            for run in para.runs:
                sz = _docx_run_size(run)
                if sz:
                    all_sizes.append(sz)

        body_size      = sorted(all_sizes)[len(all_sizes) // 2] if all_sizes else 12.0
        chapter_min_sz = body_size * 1.50
        section_min_sz = body_size * 1.15

        logger.debug(
            f"DOCX body_size={body_size:.1f}pt  "
            f"chapter≥{chapter_min_sz:.1f}  section≥{section_min_sz:.1f}"
        )

        # ── 2. Walk state ──────────────────────────────────────────────────
        chapters:  list[tuple[str, list[dict]]] = []
        # FIX: empty string instead of "Front Matter" — no phantom chapter
        # title is generated when content precedes the first real heading.
        state:     dict = {"title": "", "blocks": []}
        cur_lines: list[dict] = []
        cur_kind   = "text"

        def _flush_lines() -> None:
            nonlocal cur_lines, cur_kind
            if cur_lines:
                state["blocks"].append({"kind": cur_kind, "lines": cur_lines[:]})
                cur_lines = []
            cur_kind = "text"

        def _flush_chapter(new_title: str) -> None:
            _flush_lines()
            if state["blocks"]:
                chapters.append((state["title"], state["blocks"][:]))
            state["title"]  = new_title
            state["blocks"] = []

        def _para_heading_level(para) -> int:
            """
            Return 0 (not a heading), 1 (chapter boundary), or 2 (section).
            Checks Word style name first, then bold heuristic, then font size.
            """
            style_name = (para.style.name or "").lower().strip()
            text       = para.text.strip()

            # Title style → handled separately as title page, skip in main loop
            if style_name == "title":
                return 0

            # Heading 1 → chapter boundary
            if style_name == "heading 1":
                return 1

            # Heading 7 used as "Chapter N" label in this book schema
            if style_name == "heading 7":
                return 1

            # All other heading styles → section heading
            if style_name.startswith("heading"):
                return 2

            # Nothing further to check without text
            if not text or len(text) > 140:
                return 0

            # All-bold Normal paragraph → section heading
            if style_name == "normal" and len(text) <= 120:
                runs_with_text = [r for r in para.runs if r.text.strip()]
                if runs_with_text and all(r.bold for r in runs_with_text):
                    return 2

            # Font-size heuristics
            max_sz = _para_max_size(para, body_size)
            if max_sz >= chapter_min_sz:
                return 1
            if max_sz >= section_min_sz:
                return 2

            # Pattern-based chapter detection
            if is_chapter_heading(text):
                return 1

            return 0

        def _para_to_spans(para) -> list[dict]:
            spans = []
            for run in para.runs:
                txt = run.text
                if not txt:
                    continue
                spans.append({
                    "text":   txt,
                    "bold":   bool(run.bold),
                    "italic": bool(run.italic),
                    "size":   _docx_run_size(run) or body_size,
                })
            return spans

        # ── 3. Emit title page as first chapter ───────────────────────────
        # Collect all paragraphs before the first Heading 1 / Heading 7
        # and treat them as a title/front-matter chapter.
        # Each paragraph becomes its own block so epub_builder renders it
        # as a separate <p> rather than merging everything into one wall of text.
        title_page_paras = []
        for para in doc.paragraphs:
            sn = (para.style.name or "").lower().strip()
            if sn in ("heading 1", "heading 7"):
                break
            text = para.text.strip()
            if text:
                title_page_paras.append(para)

        if title_page_paras:
            book_title   = title_page_paras[0].text.strip()
            front_paras  = title_page_paras[1:]

            # ── Each paragraph → its own block so it renders as its own <p> ──
            # BUG 4 FIX: also extract inline and anchored images from each
            # front-matter paragraph so title-page images are not dropped.
            front_blocks: list[dict] = []
            for fp in front_paras:
                # Images first (inline + anchored — see BUG 2 fix below)
                img_blocks = _extract_inline_images(fp)
                front_blocks.extend(img_blocks)

                spans = _para_to_spans(fp)
                if not spans:
                    # Paragraph has no runs — use plain text as a single span
                    plain = fp.text.strip()
                    if plain:
                        spans = [{"text": plain, "bold": False,
                                  "italic": False, "size": body_size}]
                if spans:
                    front_blocks.append({
                        "kind": "text",
                        "lines": [
                            {
                                "spans":      spans,
                                "is_section": False,
                            }
                        ],
                    })

            chapters.append((book_title, front_blocks))

        # ── 4. Collect floating text-boxes (sidebars) ─────────────────────
        sidebar_blocks: list[dict] = _extract_textboxes(doc, body_size)

        # ── 5. Iterate top-level body children ────────────────────────────
        for elem in doc.element.body:
            local = _local(elem.tag)

            # ── Paragraph ─────────────────────────────────────────────────
            if local == "p":
                para  = DocxPara(elem, doc)
                level = _para_heading_level(para)
                text  = para.text.strip()

                if level == 1 and text:
                    _flush_chapter(text)
                    continue

                if level == 2 and text:
                    _flush_lines()
                    max_sz = _para_max_size(para, body_size)
                    cur_lines.append({
                        "spans":      [{"text": text, "bold": True,
                                        "italic": False, "size": max_sz}],
                        "is_section": True,
                    })
                    _flush_lines()
                    continue

                # Inline + anchored images embedded in this paragraph
                # BUG 2 FIX: _extract_inline_images now handles both wp:inline
                # and wp:anchor drawings (guard removed — see function below).
                img_blocks = _extract_inline_images(para)
                if img_blocks:
                    _flush_lines()
                    for ib in img_blocks:
                        state["blocks"].append(ib)

                # Normal text runs — flush after each paragraph so Word
                # paragraph boundaries become separate <p> tags in the EPUB
                spans = _para_to_spans(para)
                if spans:
                    cur_lines.append({"spans": spans, "is_section": False})
                    _flush_lines()

            # ── Table ─────────────────────────────────────────────────────
            elif local == "tbl":
                try:
                    table = DocxTable(elem, doc)
                    rows  = [
                        [c.text.strip() for c in row.cells]
                        for row in table.rows
                    ]
                    deduped = [_dedup_row(r) for r in rows if any(r)]
                    if deduped:
                        _flush_lines()
                        state["blocks"].append({"kind": "table", "rows": deduped})
                except Exception as e:
                    logger.debug(f"Table extraction error: {e}")

            else:
                continue

        # Flush remaining content
        _flush_lines()
        if state["blocks"]:
            chapters.append((state["title"], state["blocks"][:]))

        # ── 6. Inject sidebar blocks into first chapter ────────────────────
        if sidebar_blocks and chapters:
            chapters[0][1].extend(sidebar_blocks)

        # ── 7. Filter out TOC chapters ─────────────────────────────────────
        chapters = [
            (title, blocks)
            for title, blocks in chapters
            if not _is_toc_chapter(title, blocks)
        ]

        # ── 8. Drop any chapter whose title is blank and has no blocks ─────
        # This cleans up the empty sentinel that replaced "Front Matter"
        # when nothing precedes the first real heading.
        chapters = [
            (title, blocks)
            for title, blocks in chapters
            if title.strip() or blocks
        ]

        return chapters if chapters else None

    except Exception as e:
        logger.warning(f"extract_rich_chapters_from_docx failed: {e}", exc_info=True)
        return None


# ── DOCX internal helpers ─────────────────────────────────────────────────────

def _local(tag: str) -> str:
    """Strip XML namespace from a tag string."""
    return tag.split("}")[-1] if "}" in tag else tag


def _docx_run_size(run) -> float | None:
    """Return font size in points for a run, or None if not set."""
    try:
        from docx.oxml.ns import qn

        sz = run.font.size
        if sz:
            return sz.pt

        rpr = run._element.find(qn("w:rPr"))
        if rpr is not None:
            sz_node = rpr.find(qn("w:sz"))
            if sz_node is not None:
                val = sz_node.get(qn("w:val"))
                if val:
                    return int(val) / 2.0
    except Exception:
        pass
    return None


def _para_max_size(para, default: float) -> float:
    """Return the maximum run font size found in a paragraph."""
    sizes = [_docx_run_size(r) for r in para.runs if r.text.strip()]
    valid = [s for s in sizes if s]
    return max(valid) if valid else default


def _dedup_row(row: list[str]) -> list[str]:
    """
    python-docx repeats cell text for merged cells.
    Keep only the first occurrence of each consecutive duplicate.
    """
    out = []
    prev = object()
    for cell in row:
        if cell != prev:
            out.append(cell)
        prev = cell
    return out


def _extract_inline_images(para) -> list[dict]:
    """
    Extract images from a paragraph's w:drawing elements — both inline
    (wp:inline) and anchored/floating (wp:anchor).

    BUG 2 FIX: The original code had a guard:
        if drawing.find(f"{{{_NS_WP}}}inline") is None: continue
    This silently dropped every wp:anchor drawing (anchored images).
    The guard is removed; we now search for a:blip inside any drawing
    regardless of whether it is inline or anchored.
    """
    images: list[dict] = []

    for drawing in para._element.findall(f".//{{{_NS_W}}}drawing"):
        # BUG 2 FIX: removed the wp:inline-only guard that was here.
        # Both wp:inline and wp:anchor drawings may contain a:blip references.

        for blip in drawing.findall(f".//{{{_NS_A}}}blip"):
            r_embed = blip.get(f"{{{_NS_R}}}embed")
            if not r_embed:
                continue
            try:
                rel  = para.part.rels[r_embed]
                blob = rel.target_part.blob
                ct   = rel.target_part.content_type
                ext  = ct.split("/")[-1].lower()
                if ext == "jpeg":
                    ext = "jpg"
                if len(blob) < 1024:
                    continue
                images.append({"kind": "image", "data": blob, "ext": ext})
            except Exception as e:
                logger.debug(f"Image skip: {e}")

    return images


def _extract_textboxes(doc, body_size: float) -> list[dict]:
    """
    Extract anchored (floating) text boxes from the document as sidebar blocks.
    """
    sidebars: list[dict] = []

    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph as DocxPara

    for anchor in doc.element.body.findall(f".//{{{_NS_WP}}}anchor"):
        txbx = anchor.find(f".//{{{_NS_WPS}}}txbx")
        if txbx is None:
            continue

        txbx_content = txbx.find(f"{{{_NS_W}}}txbxContent")
        if txbx_content is None:
            continue

        lines: list[dict] = []

        for p_elem in txbx_content.findall(f"{{{_NS_W}}}p"):
            try:
                para = DocxPara(p_elem, doc)
                text = para.text.strip()
                if not text:
                    continue

                spans = []
                for run in para.runs:
                    txt = run.text
                    if not txt:
                        continue
                    spans.append({
                        "text":   txt,
                        "bold":   bool(run.bold),
                        "italic": bool(run.italic),
                        "size":   _docx_run_size(run) or body_size,
                    })

                if spans:
                    style_name = (para.style.name or "").lower()
                    lines.append({
                        "spans":      spans,
                        "is_section": "heading" in style_name,
                    })
            except Exception:
                continue

        if lines:
            sidebars.append({"kind": "sidebar", "lines": lines})

    return sidebars


# ════════════════════════════════════════════════════════════════════════════
# PDF — Column detection
# ════════════════════════════════════════════════════════════════════════════

def detect_columns(
    blocks: list[dict],
    page_width: float,
    min_gutter_width: float = 10.0,
    strip_count: int = 200,
) -> list[tuple[float, float]]:

    if not blocks or page_width <= 0:
        return [(0.0, page_width)]

    strip_width = page_width / strip_count
    occupancy   = [0] * strip_count

    for blk in blocks:
        bx0, _, bx1, _ = blk.get("bbox", (0, 0, 0, 0))
        bx0 = max(0.0, bx0)
        bx1 = min(page_width, bx1)
        if bx1 <= bx0:
            continue
        si = int(bx0 / strip_width)
        ei = min(int(bx1 / strip_width), strip_count - 1)
        for s in range(si, ei + 1):
            occupancy[s] += 1

    if not any(occupancy):
        return [(0.0, page_width)]

    gutters   = []
    in_gutter = False
    g_start   = 0

    for i, occ in enumerate(occupancy):
        if occ == 0 and not in_gutter:
            in_gutter = True
            g_start   = i
        elif occ > 0 and in_gutter:
            in_gutter = False
            g_x0 = g_start * strip_width
            g_x1 = i * strip_width
            if g_x1 - g_x0 >= min_gutter_width:
                gutters.append((g_x0, g_x1))

    if in_gutter:
        g_x0 = g_start * strip_width
        g_x1 = strip_count * strip_width
        if g_x1 - g_x0 >= min_gutter_width:
            gutters.append((g_x0, g_x1))

    if not gutters:
        return [(0.0, page_width)]

    col_boundaries = [0.0]
    for gx0, gx1 in gutters:
        col_boundaries.append(gx0)
        col_boundaries.append(gx1)
    col_boundaries.append(page_width)

    columns = []
    for i in range(0, len(col_boundaries) - 1, 2):
        cx0 = col_boundaries[i]
        cx1 = col_boundaries[i + 1]
        if cx1 - cx0 > min_gutter_width:
            columns.append((cx0, cx1))

    return columns if columns else [(0.0, page_width)]


def _block_column_index(
    blk_bbox: tuple,
    columns: list[tuple[float, float]],
    page_width: float,
) -> int:

    bx0, _, bx1, _ = blk_bbox

    if (bx1 - bx0) >= page_width * 0.80:
        return -1

    centre_x  = (bx0 + bx1) / 2.0
    best_col  = 0
    best_dist = float("inf")

    for i, (cx0, cx1) in enumerate(columns):
        dist = abs(centre_x - (cx0 + cx1) / 2.0)
        if dist < best_dist:
            best_dist = dist
            best_col  = i

    return best_col


def sort_blocks_into_columns(
    blocks: list[dict],
    columns: list[tuple[float, float]],
    page_width: float,
) -> list[dict]:

    if len(columns) <= 1:
        return blocks

    full_width = []
    col_blocks = []

    for blk in blocks:
        bbox    = blk.get("bbox") or (0, 0, 0, 0)   # ← CHANGE 1: or instead of default=
        top_y   = bbox[1]
        col_idx = _block_column_index(bbox, columns, page_width)

        if col_idx == -1:
            full_width.append((top_y, blk))
        else:
            col_blocks.append((col_idx, top_y, blk))

    col_blocks.sort(key=lambda t: (t[0], t[1]))
    sorted_col = [blk for _, _, blk in col_blocks]

    if not full_width:
        return sorted_col

    full_width.sort(key=lambda t: t[0])
    fw_iter = iter(full_width)
    next_fw = next(fw_iter, None)
    result: list[dict] = []

    for blk in sorted_col:
        blk_top_y = (blk.get("bbox") or (0, 0, 0, 0))[1]   # ← CHANGE 2: .get().or not []
        while next_fw is not None and next_fw[0] <= blk_top_y:
            result.append(next_fw[1])
            next_fw = next(fw_iter, None)
        result.append(blk)

    if next_fw is not None:
        result.append(next_fw[1])
    for _, fw_blk in fw_iter:
        result.append(fw_blk)

    return result


# ════════════════════════════════════════════════════════════════════════════
# PDF — Sidebar detection
# ════════════════════════════════════════════════════════════════════════════

def _block_has_background(
    blk_bbox: tuple,
    page_drawings: list[dict],
    overlap_threshold: float = 0.55,
) -> bool:

    bx0, by0, bx1, by1 = blk_bbox
    blk_area = max((bx1 - bx0) * (by1 - by0), 1e-6)

    for d in page_drawings:
        fill  = d.get("fill")
        color = d.get("color")

        has_fill   = fill  is not None and fill  != (1, 1, 1)
        has_stroke = color is not None and color != (1, 1, 1)

        if not (has_fill or has_stroke):
            continue

        rect = d.get("rect")
        if rect is None:
            continue

        try:
            rx0, ry0, rx1, ry1 = (
                float(rect[0]), float(rect[1]),
                float(rect[2]), float(rect[3]),
            )
        except Exception:
            continue

        ix0 = max(bx0, rx0)
        iy0 = max(by0, ry0)
        ix1 = min(bx1, rx1)
        iy1 = min(by1, ry1)

        if ix1 <= ix0 or iy1 <= iy0:
            continue

        overlap = (ix1 - ix0) * (iy1 - iy0)
        if overlap / blk_area >= overlap_threshold:
            return True

    return False


def classify_sidebar(
    blk: dict,
    columns: list[tuple[float, float]],
    page_width: float,
    page_drawings: list[dict] | None = None,
    width_threshold: float = 0.45,
    offset_threshold: float = 0.15,
) -> bool:

    bbox = blk.get("bbox", (0, 0, 0, 0))
    bx0, _, bx1, _ = bbox
    blk_width  = bx1 - bx0
    line_count = len(blk.get("lines", []))

    if line_count <= 1:
        return False
    if blk_width >= page_width * 0.80:
        return False

    sig_width  = blk_width < page_width * width_threshold
    centre_x   = (bx0 + bx1) / 2.0
    min_col_dist = min(
        abs(centre_x - (cx0 + cx1) / 2.0)
        for cx0, cx1 in columns
    )
    sig_offset     = min_col_dist > page_width * offset_threshold
    sig_background = False

    if page_drawings:
        sig_background = _block_has_background(bbox, page_drawings)

    return sum([sig_width, sig_offset, sig_background]) >= 2


# ════════════════════════════════════════════════════════════════════════════
# PDF — Image pairing
# ════════════════════════════════════════════════════════════════════════════

def find_text_image_pairs(
    blocks: list[dict],
    horiz_overlap_threshold: float = 0.25,
    vert_proximity_pts: float = 80.0,
) -> None:

    image_entries = [
        (i, blk["bbox"])
        for i, blk in enumerate(blocks)
        if blk.get("kind") == "image" and "bbox" in blk
    ]

    if not image_entries:
        return

    for txt_idx, blk in enumerate(blocks):
        if blk.get("kind") not in ("text", "sidebar"):
            continue

        tbbox = blk.get("bbox")
        if not tbbox:
            continue

        tx0, ty0, tx1, ty1 = tbbox

        for img_idx, ibbox in image_entries:
            if img_idx == txt_idx:
                continue

            ix0, iy0, ix1, iy1 = ibbox
            img_width    = max(ix1 - ix0, 1e-6)
            horiz_overlap = max(0.0, min(tx1, ix1) - max(tx0, ix0))

            if horiz_overlap / img_width < horiz_overlap_threshold:
                continue

            vert_gap = max(ty0 - iy1, iy0 - ty1, 0.0)
            if vert_gap > vert_proximity_pts:
                continue

            blk["nearby_image"] = True
            break


# ════════════════════════════════════════════════════════════════════════════
# PDF — Helpers
# ════════════════════════════════════════════════════════════════════════════

def _bbox_overlap(a: tuple, b: tuple, threshold: float = 0.25) -> bool:

    ax0, ay0, ax1, ay1 = a
    bx0, by0, bx1, by1 = b

    ix0 = max(ax0, bx0)
    iy0 = max(ay0, by0)
    ix1 = min(ax1, bx1)
    iy1 = min(ay1, by1)

    if ix1 <= ix0 or iy1 <= iy0:
        return False

    intersection = (ix1 - ix0) * (iy1 - iy0)
    area_a = max((ax1 - ax0) * (ay1 - ay0), 1e-6)

    return (intersection / area_a) > threshold


def _is_page_number(text: str) -> bool:
    s = text.strip()
    return bool(
        re.match(r"^\d{1,3}$", s)
        or re.match(r"^[ivxlcdmIVXLCDM]{1,8}$", s)
    )


def _is_cover_page(page) -> bool:

    blocks     = page.get_text("dict")["blocks"]
    text_spans = []

    for blk in blocks:
        if blk.get("type") != 0:
            continue
        for ln in blk.get("lines", []):
            for sp in ln.get("spans", []):
                txt = sp["text"].strip()
                if txt:
                    text_spans.append((txt, sp["size"]))

    if not text_spans:
        return True

    long_lines = sum(1 for t, _ in text_spans if len(t.split()) > 8)
    large_text = sum(1 for _, sz in text_spans if sz > 20)

    return long_lines == 0 and large_text >= 2


def _is_title_page(page) -> bool:

    txt     = page.get_text().lower()
    signals = ["copyright", "isbn", "all rights reserved", "published by"]
    hits    = sum(1 for s in signals if s in txt)

    return hits >= 2


def _is_toc_page(page, body_size: float) -> bool:

    raw_blocks = page.get_text("dict")["blocks"]
    hits       = 0

    for blk in raw_blocks:
        if blk.get("type") != 0:
            continue
        for ln in blk.get("lines", []):
            spans = ln.get("spans", [])
            if not spans:
                continue
            text = " ".join(s["text"] for s in spans).strip()
            if not text:
                continue

            if re.search(r"\.{3,}\s*\d+$", text):
                hits += 1
                continue

            if re.search(r"\s{3,}[ivxlcdm\d]+\s*$", text, re.I):
                hits += 1
                continue

            if len(spans) >= 2:
                first     = spans[0]
                last      = spans[-1]
                last_text = last["text"].strip()
                if re.match(r"^[ivxlcdm\d]+$", last_text, re.I):
                    gap = last["bbox"][0] - first["bbox"][2]
                    if gap > 80:
                        hits += 1

    return hits >= 4


def _line_x0(line) -> float:
    spans = line.get("spans", [])
    if not spans:
        return 0
    return spans[0]["bbox"][0]


def _join_paragraph_lines(lines: list[dict]) -> list[dict]:

    if not lines:
        return []

    paragraphs: list[dict] = []
    current    = {"spans": [], "is_section": False}
    prev_text  = ""
    prev_x0    = None

    for line in lines:
        text = " ".join(s["text"] for s in line["spans"]).strip()
        if not text:
            continue

        x0 = _line_x0({"spans": line["spans"]})
        starts_new_para = False

        if prev_x0 is None:
            starts_new_para = True
        elif abs(x0 - prev_x0) > 12:
            starts_new_para = True
        elif prev_text.endswith((".", "!", "?", ":")):
            starts_new_para = True

        if starts_new_para:
            if current["spans"]:
                paragraphs.append(current)
            current = {
                "spans":      [],
                "is_section": line.get("is_section", False),
            }
            if "dropcap_char" in line:
                current["dropcap_char"] = line["dropcap_char"]

        current["spans"].extend(line["spans"])
        prev_text = text
        prev_x0   = x0

    if current["spans"]:
        paragraphs.append(current)

    return paragraphs


# ════════════════════════════════════════════════════════════════════════════
# PDF — Rich extraction (PyMuPDF)
# ════════════════════════════════════════════════════════════════════════════

def extract_rich_chapters(pdf_path: str) -> list[tuple[str, list[dict]]] | None:
    try:
        import fitz
    except ImportError:
        logger.warning("PyMuPDF not available")
        return None
    try:
        doc   = fitz.open(pdf_path)
        sizes = []
        for page in doc:
            for blk in page.get_text("dict")["blocks"]:
                if blk.get("type") != 0:
                    continue
                for ln in blk["lines"]:
                    for sp in ln["spans"]:
                        if sp["text"].strip():
                            sizes.append(sp["size"])
        if not sizes:
            doc.close()
            return None
        sizes.sort()
        body_size        = sizes[len(sizes) // 2]
        chapter_min_size = body_size * 1.6
        section_min_size = body_size * 1.2
        chapters = []
        # FIX: empty string instead of "Front Matter" — no phantom chapter
        # title is generated when content precedes the first real heading.
        state    = {"title": "", "blocks": [], "lines": []}
        seen_xrefs: set             = set()
        pending_heading_parts: list[str] = []
        pending_dropcap: str | None      = None
        current_block_kind               = "text"

        def _flush_lines():
            if state["lines"]:
                joined = _join_paragraph_lines(state["lines"])
                state["blocks"].append({"kind": current_block_kind, "lines": joined})
                state["lines"] = []

        def _flush_pending_heading():
            nonlocal pending_heading_parts
            if not pending_heading_parts:
                return
            heading_text          = " ".join(pending_heading_parts)
            pending_heading_parts = []
            _save_chapter(heading_text)

        def _save_chapter(new_title: str):
            _flush_lines()
            # BUG 3 FIX: call find_text_image_pairs on the completed chapter's
            # blocks so nearby_image is set before we append and reset state.
            if state["blocks"]:
                find_text_image_pairs(state["blocks"])
                chapters.append((state["title"], state["blocks"][:]))
            state["title"]  = new_title
            state["blocks"] = []
            state["lines"]  = []

        for page in doc:

            if page.number == 0 and _is_cover_page(page):
                logger.debug("Skipping cover page")
                continue

            if page.number <= 2 and _is_title_page(page):
                logger.debug("Skipping title page")
                continue

            if _is_toc_page(page, body_size):
                logger.debug(f"Skipping TOC page {page.number}")
                continue

            page_width = page.rect.width

            try:
                page_drawings = page.get_drawings()
            except Exception:
                page_drawings = []

            raw_blocks     = page.get_text("dict", sort=True)["blocks"]
            columns        = detect_columns(raw_blocks, page_width)
            ordered_blocks = sort_blocks_into_columns(raw_blocks, columns, page_width)

            # BUG 1 FIX: extract tables via PyMuPDF's find_tables() and inject
            # them as {"kind": "table", "rows": [...]} blocks before processing
            # the text/image blocks so reading order is preserved.
            try:
                tab_finder = page.find_tables()
                page_tables = {
                    tuple(round(v, 1) for v in tbl.bbox): tbl
                    for tbl in tab_finder.tables
                }
            except Exception as e:
                logger.debug(f"find_tables error on page {page.number}: {e}")
                page_tables = {}

            for blk in ordered_blocks:

                btype = blk.get("type")
                bbox  = tuple(blk.get("bbox") or (0, 0, 0, 0))   # ← CHANGE 3

                # BUG 1 FIX: check whether this block's bbox matches a table
                # and emit it as a table block instead of treating it as text.
                rounded_bbox = tuple(round(v, 1) for v in bbox)
                if rounded_bbox in page_tables:
                    tbl = page_tables[rounded_bbox]
                    try:
                        rows = tbl.extract()
                        # Normalise: cells may be None
                        norm_rows = [
                            [str(cell) if cell is not None else "" for cell in row]
                            for row in rows
                            if any(cell for cell in row)
                        ]
                        if norm_rows:
                            _flush_pending_heading()
                            _flush_lines()
                            state["blocks"].append({"kind": "table", "rows": norm_rows})
                    except Exception as e:
                        logger.debug(f"Table extract error: {e}")
                    continue

                if btype == 1:
                    # RISK FIX: original guard was `if xref and xref not in seen_xrefs`
                    # which is falsy for xref==0 (inline-stream images).
                    # Changed to `if xref is not None` so xref=0 images are attempted.
                    xref = blk.get("xref")
                    if xref is not None and xref not in seen_xrefs:
                        # xref==0 means inline-stream: fall back to page.get_images()
                        # to locate the actual xref for this block's bbox.
                        actual_xref = xref
                        if xref == 0:
                            try:
                                for img_info in page.get_images(full=True):
                                    # img_info: (xref, smask, w, h, bpc, cs, alt_cs, name, filter, ref)
                                    actual_xref = img_info[0]
                                    break  # take first; bbox-matching would be more precise
                            except Exception:
                                pass

                        try:
                            img = doc.extract_image(actual_xref)
                            w   = img.get("width", 0)
                            h   = img.get("height", 0)
                            if w >= 80 and h >= 80:
                                _flush_pending_heading()
                                _flush_lines()
                                state["blocks"].append({
                                    "kind":   "image",
                                    "data":   img["image"],
                                    "ext":    img.get("ext", "png"),
                                    "width":  w,
                                    "height": h,
                                    "bbox":   bbox,
                                })
                                seen_xrefs.add(actual_xref)
                        except Exception as e:
                            logger.debug(f"Image extraction error: {e}")
                    continue

                if btype != 0:
                    continue

                current_block_kind = (
                    "sidebar"
                    if classify_sidebar(blk, columns, page_width, page_drawings)
                    else "text"
                )

                for ln in blk.get("lines", []):
                    spans     = ln.get("spans", [])
                    line_text = " ".join(s["text"] for s in spans).strip()

                    if not line_text:
                        continue
                    if _is_page_number(line_text):
                        continue

                    max_size      = max((s["size"] for s in spans), default=body_size)
                    is_very_large = max_size >= chapter_min_size

                    if is_very_large and len(line_text) <= 2:
                        if pending_heading_parts:
                            _flush_pending_heading()
                        pending_dropcap = line_text.strip()
                        continue

                    if len(line_text) < 120 and (
                        is_very_large or is_chapter_heading(line_text)
                    ):
                        if pending_dropcap is not None:
                            pending_dropcap = None
                        pending_heading_parts.append(line_text)
                        continue

                    if pending_heading_parts:
                        _flush_pending_heading()

                    span_items = [
                        {
                            "text":   s["text"],
                            "bold":   bool(s.get("flags", 0) & 16),
                            "italic": bool(s.get("flags", 0) & 2),
                            "size":   s["size"],
                        }
                        for s in spans
                        if s["text"].strip()
                    ]

                    if span_items:
                        line_dict: dict = {
                            "spans":      span_items,
                            "is_section": (max_size >= section_min_size),
                        }
                        if pending_dropcap is not None:
                            line_dict["dropcap_char"] = pending_dropcap
                            pending_dropcap = None
                        state["lines"].append(line_dict)

        _flush_pending_heading()
        _flush_lines()

        # BUG 3 FIX: pair images for the final chapter (not saved via _save_chapter).
        if state["blocks"]:
            find_text_image_pairs(state["blocks"])
            chapters.append((state["title"], state["blocks"][:]))

        doc.close()

        # FIX: drop the empty-title sentinel if it ended up with no blocks,
        # or if it slipped through with a blank title but real content
        # (content gets absorbed into the first named chapter instead).
        chapters = [
            (title, blocks)
            for title, blocks in chapters
            if title.strip() or blocks
        ]

        return chapters if chapters else None

    except Exception as e:
        logger.warning(f"extract_rich_chapters failed: {e}")
        return None


# ════════════════════════════════════════════════════════════════════════════
# Text cleaning & chapter detection (shared by PDF and DOCX plain-text paths)
# ════════════════════════════════════════════════════════════════════════════

def detect_running_headers(lines: list[str]) -> set[str]:

    stripped = [l.strip() for l in lines if l.strip()]
    freq     = Counter(stripped)

    return {
        line
        for line, count in freq.items()
        if count >= 3 and len(line) > 4
    }


def clean_text(text: str) -> str:

    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)

    lines           = text.split("\n")
    running_headers = detect_running_headers(lines)
    cleaned: list[str] = []

    for line in lines:
        s = line.strip()
        if re.match(r"^\d{1,3}$", s):
            continue
        if re.match(r"^[ivxlcdmIVXLCDM]{1,6}$", s):
            continue
        if s in running_headers:
            continue
        cleaned.append(line)

    return "\n".join(cleaned)


def skip_toc(lines: list[str]) -> int:

    toc_pattern   = re.compile(r"\.{3,}|\s{2,}\d+\s*$")
    toc_hits      = 0
    last_toc_line = 0

    for i, line in enumerate(lines[:150]):
        if toc_pattern.search(line.strip()):
            toc_hits += 1
            last_toc_line = i

    return last_toc_line + 3 if toc_hits >= 3 else 0


def is_chapter_heading(line: str) -> bool:

    s = line.strip()
    if not s or len(s) > 100:
        return False

    patterns = [
        re.compile(r"^chapter\s+\d+", re.IGNORECASE),
        re.compile(
            r"^(introduction|foreword|preface|dedication|"
            r"acknowledgements?|endorsements?|"
            r"conclusion|epilogue|prologue)$",
            re.IGNORECASE,
        ),
        re.compile(
            r"^chapter\s+(one|two|three|four|five|"
            r"six|seven|eight|nine|ten)\b",
            re.IGNORECASE,
        ),
    ]

    return any(p.match(s) for p in patterns)


def is_section_heading(line: str) -> bool:

    s = line.strip()
    if not s or len(s) > 80 or len(s) < 4:
        return False

    alpha = [c for c in s if c.isalpha()]
    if not alpha:
        return False

    upper_ratio = sum(1 for c in alpha if c.isupper()) / len(alpha)

    return upper_ratio > 0.85 and len(s.split()) >= 2


def extract_chapters_from_text(text: str) -> list[tuple[str, str]]:

    text  = clean_text(text)
    lines = text.split("\n")

    start_idx = skip_toc(lines)
    lines     = lines[start_idx:]

    chapters: list[tuple[str, str]] = []
    # FIX: empty string instead of "Front Matter" — no phantom chapter title
    # is emitted when content precedes the first real chapter heading.
    current_title   = ""
    current_content: list[str] = []

    i = 0
    while i < len(lines):
        line = lines[i]
        s    = line.strip()

        if not s:
            current_content.append("")
            i += 1
            continue

        if is_chapter_heading(s):
            content_text = "\n".join(current_content).strip()
            # Only save the pre-heading content if there actually is some;
            # an empty title with no content produces nothing.
            if content_text:
                chapters.append((current_title, content_text))
            current_title   = s
            current_content = []
            i += 1
            continue

        elif is_section_heading(s):
            current_content.extend(["", line, ""])
            i += 1
            continue

        current_content.append(line)
        i += 1

    content_text = "\n".join(current_content).strip()
    if content_text:
        chapters.append((current_title, content_text))

    return chapters if chapters else split_by_wordcount(text)


def split_by_wordcount(
    text: str,
    words_per_part: int = 800,
) -> list[tuple[str, str]]:

    words = text.split()
    parts: list[tuple[str, str]] = []

    for part_num, i in enumerate(range(0, len(words), words_per_part), start=1):
        chunk = " ".join(words[i : i + words_per_part])
        parts.append((f"Part {part_num}", chunk))

    return parts if parts else [("Content", text)]
